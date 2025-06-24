from flask import Flask, render_template, request, session, redirect, url_for,send_file
from flask import send_file
from docx import Document
import io

from google import genai
import google.generativeai as genai



app = Flask(__name__)



@app.route('/')
def index():
    return render_template('index.html')



@app.route('/submit', methods=['POST'])
def submit():
    path = request.form.get('path')
    course_name = request.form.get('courseName')
    course_description = request.form.get('courseDescription')

    highlights = request.form.getlist('highlights[]')
    topics_titles = request.form.getlist('topics_title[]')
    topics_descs = request.form.getlist('topics_desc[]')
    outcomes = [value for key, value in request.form.items() if key.startswith('outcome')]
    activities = request.form.getlist('activities[]')

    # Create Word document
    doc = Document()
    doc.add_heading('progress report', 0)

    doc.add_heading('Course Path', level=1)
    doc.add_paragraph(path)

    doc.add_heading('Course Name', level=1)
    doc.add_paragraph(course_name)

    doc.add_heading('Course Description', level=1)
    doc.add_paragraph(course_description)

    doc.add_heading('Highlights', level=1)
    for h in highlights:
        doc.add_paragraph(f'{h}', style='List Bullet')

    doc.add_heading('Topics', level=1)
    for title, desc in zip(topics_titles, topics_descs):
        p = doc.add_paragraph(style='List Bullet')  # bullet point
        run_title = p.add_run(f"{title}: ")
        run_title.bold = True
        p.add_run(desc)

    doc.add_heading('Outcomes', level=1)
    for o in outcomes:
        doc.add_paragraph(f'{o}', style='List Bullet')

    doc.add_heading('Activities', level=1)
    for a in activities:
        doc.add_paragraph(f'{a}', style='List Bullet')

    # ✅ Save to BytesIO and send
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    return send_file(
        doc_io,
        as_attachment=True,
        download_name=f'{path}-progress report.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

# @app.route('/submit', methods=['POST'])
# def submit():
#     path = request.form.get('path')
#     course_name = request.form.get('courseName')
#     course_description = request.form.get('courseDescription')

#     highlights = request.form.getlist('highlights[]')
#     topics_titles = request.form.getlist('topics_title[]')
#     topics_descs = request.form.getlist('topics_desc[]')
#     outcomes = [value for key, value in request.form.items() if key.startswith('outcome')]
#     activities = request.form.getlist('activities[]')
            
#     return render_template('download.html', path=path,
#                                             course_name=course_name,
#                                             course_description=course_description,
#                                             highlights=highlights,
#                                             topics=list(zip(topics_titles, topics_descs)),
#                                             outcomes=outcomes,
#                                             activities=activities)

if __name__ == '__main__':
    app.run()